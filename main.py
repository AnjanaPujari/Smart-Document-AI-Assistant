from fastapi import FastAPI, UploadFile, File
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse

import faiss
import uuid
import os
from pypdf import PdfReader
from docx import Document
from langchain_groq import ChatGroq


os.environ["TOKENIZERS_PARALLELISM"] = "false"
os.environ["OMP_NUM_THREADS"] = "1"
os.environ["MKL_NUM_THREADS"] = "1"


# MODELS

llm = ChatGroq(
    model="openai/gpt-oss-20b",
    temperature=0
)


# FAISS

index = faiss.IndexFlatIP(384)

chunks = []


# EMBEDDING MODEL

def create_embedding_model():

    from sentence_transformers import SentenceTransformer

    return SentenceTransformer(
        "all-MiniLM-L6-v2",
        device="cpu"
    )


# SEARCH DOCUMENTS

def search_documents(query, chunks, k=3):

    embedding_model = create_embedding_model()

    query_embedding = embedding_model.encode(
        [query],
        normalize_embeddings=True
    )

    del embedding_model

    distances, indices = index.search(
        query_embedding,
        k
    )

    valid_chunks = []

    for i in indices[0]:

        if i >= 0 and i < len(chunks):

            valid_chunks.append(
                chunks[i]
            )

    return valid_chunks, distances


# FASTAPI

app = FastAPI()


# FRONTEND

app.mount(
    "/static",
    StaticFiles(directory="frontend"),
    name="static"
)


@app.get("/app")
def serve_frontend():

    return FileResponse(
        "frontend/index.html"
    )


# BASIC ENDPOINTS

@app.get("/")
def home():

    return {
        "message":
            "Hello, Smart Document AI Assistant!"
    }


@app.get("/status")
def status():

    return {
        "status":
            "Server is running successfully"
    }


# EXTRACT TEXT FROM PDF

def extract_pdf_text(file_path):

    reader = PdfReader(file_path)

    pages = len(reader.pages)

    all_chunks = []

    chunk_size = 250
    overlap = 50

    for page in reader.pages:

        page_text = page.extract_text() or ""

        page_text = page_text.strip()

        for i in range(
            0,
            len(page_text),
            chunk_size - overlap
        ):

            chunk = page_text[
                i:i + chunk_size
            ]

            if chunk.strip():

                all_chunks.append(
                    chunk.strip()
                )

    return all_chunks, pages


# EXTRACT TEXT FROM DOCX

def extract_docx_text(file_path):

    document = Document(file_path)

    full_text = []

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:

            full_text.append(text)

    for table in document.tables:

        for row in table.rows:

            row_text = []

            for cell in row.cells:

                text = cell.text.strip()

                if text:

                    row_text.append(text)

            if row_text:

                full_text.append(
                    " | ".join(row_text)
                )

    complete_text = "\n".join(
        full_text
    )

    all_chunks = []

    chunk_size = 250
    overlap = 50

    for i in range(
        0,
        len(complete_text),
        chunk_size - overlap
    ):

        chunk = complete_text[
            i:i + chunk_size
        ]

        if chunk.strip():

            all_chunks.append(
                chunk.strip()
            )

    return all_chunks


# UPLOAD PDF / DOCX

@app.post("/upload")
def upload_file(
    file: UploadFile = File(...)
):

    global chunks

    file_extension = os.path.splitext(
        file.filename
    )[1].lower()

    allowed_extensions = [
        ".pdf",
        ".docx"
    ]

    if file_extension not in allowed_extensions:

        return {
            "error":
                "Only PDF and DOCX files are allowed."
        }

    os.makedirs(
        "uploads",
        exist_ok=True
    )

    file_name = (
        f"{uuid.uuid4()}_{file.filename}"
    )

    file_path = os.path.join(
        "uploads",
        file_name
    )

    with open(
        file_path,
        "wb"
    ) as buffer:

        buffer.write(
            file.file.read()
        )

    if file_extension == ".pdf":

        all_chunks, pages = extract_pdf_text(
            file_path
        )

    else:

        all_chunks = extract_docx_text(
            file_path
        )

        pages = None

    if not all_chunks:

        return {
            "error":
                "No readable text found in the document."
        }

    chunks = all_chunks

    embedding_model = create_embedding_model()

    embeddings = embedding_model.encode(
        chunks,
        normalize_embeddings=True
    )

    del embedding_model

    index.reset()

    index.add(embeddings)

    print(
        "FAISS index size:",
        index.ntotal
    )

    return {
        "message":
            "File uploaded and text extracted successfully",

        "filename":
            file.filename,

        "file_type":
            file_extension,

        "pages":
            pages,

        "number_of_chunks":
            len(chunks)
    }


# SEARCH

@app.post("/search")
def search(query: str):

    if not chunks:

        return {
            "query":
                query,

            "answer":
                "Please upload a document before asking a question."
        }

    matching_chunks, distances = search_documents(
        query,
        chunks,
        k=3
    )

    if not matching_chunks:

        return {
            "query":
                query,

            "answer":
                "Information not found in the document."
        }

    print(
        "QUESTION:",
        query
    )

    context = "\n\n".join(
        matching_chunks
    )

    print(
        "RETRIEVED CONTEXT:"
    )

    print(context)

    answer = generate_answer(
        query,
        context
    )

    print(
        "Generated answer:",
        answer
    )

    return {
        "query":
            query,

        "matching_chunks":
            matching_chunks,

        "answer":
            answer
    }


# GROQ ANSWER GENERATION

def generate_answer(query, context):

    prompt = f"""
You are a document question-answering assistant.

Use ONLY the information provided in the context.

Context:
{context}

Question:
{query}

Answer the question directly and completely.
Do not use outside knowledge.
Do not include unrelated information.

If the answer is not present in the context, say:
Information not found in the document.

Answer:
"""

    response = llm.invoke(prompt)

    return response.content.strip()